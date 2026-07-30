# -*- coding: utf-8 -*-
"""A2b 報告產生器。

基礎設施(set_font/para/heading/add_toc/add_page_footer/table/fig/count_units)
沿用 A2a `Assessment2a/build_report.py`;內文、表格、圖表全部重寫。

所有數字來源:Huang_26254793_421104_Assessment 2b.ipynb(章節編號標於各表註)。
執行:python build_report.py
"""
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

EAST = "微軟正黑體"
HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "Huang_26254793_421104_Assessment 2b.docx")
NB = "Huang_26254793_421104_Assessment 2b.ipynb"


# ───────────────────────── 基礎設施(沿用 A2a)─────────────────────────
def set_font(run, size=11, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), EAST)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def para(doc, text, size=11, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size, bold)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def heading(doc, text):
    h = doc.add_heading(level=1)
    r = h.add_run(text)
    set_font(r, 14, True)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_toc(doc, entries):
    p = doc.add_paragraph()
    r = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin"); f1.set(qn("w:dirty"), "true")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = 'TOC \\o "1-1" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    # 欄位快取寫入靜態條目:Word 開啟會自動更新,不支援欄位的預覽器仍看得到目錄
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve")
    t.text = "　|　".join(entries)
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for el in (f1, it, f2, t, f3):
        r._r.append(el)


def add_page_footer(doc):
    fp = doc.sections[0].footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for el in (f1, it, f3):
        r._r.append(el)
    set_font(r, 9)


NOTES = []


def table(doc, caption, headers, rows, source_note, font_size=9):
    NOTES.append(source_note)
    cap = doc.add_paragraph(); rc = cap.add_run(caption); set_font(rc, 10, True)
    cap.paragraph_format.space_after = Pt(3)
    cap.paragraph_format.keep_with_next = True
    tb = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tb.style = "Table Grid"
    trPr = tb.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader"); th.set(qn("w:val"), "true"); trPr.append(th)
    for row in tb.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    for j, h in enumerate(headers):
        set_font(tb.rows[0].cells[j].paragraphs[0].add_run(h), font_size, True)
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            set_font(tb.rows[i].cells[j].paragraphs[0].add_run(str(v)), font_size)
    note = doc.add_paragraph(); set_font(note.add_run(source_note), 8.5, italic=True)
    note.paragraph_format.space_after = Pt(10)
    return tb


def fig(doc, caption, path, source_note):
    NOTES.append(source_note)
    cap = doc.add_paragraph(); rc = cap.add_run(caption); set_font(rc, 10, True)
    cap.paragraph_format.space_after = Pt(3)
    cap.paragraph_format.keep_with_next = True
    doc.add_picture(path, width=Inches(5.85))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.keep_with_next = True
    note = doc.add_paragraph(); set_font(note.add_run(source_note), 8.5, italic=True)
    note.paragraph_format.space_after = Pt(10)


def count_units(texts):
    """MS Word 近似口徑:CJK 字元逐字計 1,英數連續串計 1。"""
    n = 0
    for t in texts:
        n += len(re.findall("[一-鿿　-〿＀-￯]", t))
        n += len(re.findall("[A-Za-z0-9]+", t))
    return n


# ═══════════════════════════════ 內文 ═══════════════════════════════
BODY = {}

BODY['管理摘要'] = [
    '本報告以 K-means 將 107 位 IBM 顧問依三項績效指標分為四群,作為獎金級距依據(圖五)。「帶隊核心群」(4 人)為績效最優群,惟僅佔 3.7%;「高投入待認可群」(25 人)投入度達平均 2.2 倍卻幾乎未獲指派,未標準化時無法辨識。',

    '提請決定:級距配置比例(圖七)、是否併合最高兩級、是否檢視指派流程。',

]

BODY['一、主要數據特徵:納入與排除的決策'] = [
    '四欄的業務意義、統計特性與納入決策逐欄列於表一。保留的三項指標分別涵蓋投入程度、獲得指派的機會與承擔的責任層級,共同構成績效全貌;EmployeeID 則予排除——它是標稱型識別碼,數值不承載順序意義,且標準差為其餘三欄的 36 至 103 倍,納入後獨佔距離平方 99% 以上,實跑反證顯示分群會退化為按編號切段。',

    '第二層決策是尺度。常見誤解是「不標準化較中立」,事實相反:未縮放時權重由各欄離散程度而非業務判斷決定——認可度一欄即佔距離平方的 79.6%,使分群退化為單變數複製,與「直接按認可度分組」的調整後蘭德指數(Hubert & Arabie, 1985)達 0.99。',

    '其影響可具體量化(圖三、表二):10 位使用率 0.61 至 0.92 但認可度為 0 的顧問,未縮放時被併入一個使用率跨幅 0.00 至 0.92 的 86 人群。縮放方式本身則不影響結論:z 分數、MinMax、以及對數轉換後標準化(因認可度與領導角色右偏)三者皆得相同的四群結構與績效最優群(notebook §5.1b)。標準化並非「不加權」,而是把權重改為可依政策調整的假設(Milligan & Cooper, 1988)。',

]

BODY['二、最優聚類數的決定'] = [
    '四項判準於掃描前即已訂定,以免先看結果再找理由;門檻與逐 k 裁決見表三:統計面為輪廓係數(Rousseeuw, 1987)與肘部法,業務面為每群最小人數與群內單一 KPI 全距。',

    '統計面兩項衝突:輪廓係數最高在 k = 5(0.6823),但 k = 3、4、5 全距僅 0.0146、視為平手;肘部法指向 k = 4(圖一)。業務面則排除相鄰候選:k = 5 產生兩個各 2 人的群;k = 3 的中間群 27 人,群內認可度全距等於全體全距的 100%,該指標在此級距內無區分力,k = 4 拆開後降至 50% 以下。',

    '判準門檻亦經檢驗:取 3 或 4 人結論同為 k = 4,惟該門檻無外部依據(§2.5c)。k = 4 的種子平均 ARI 為 0.961,僅 10 人曾變動且全在邊界,個體輪廓值僅 1 人為負(圖四)。',

]

BODY['三、四個群組的語言解釋'] = [
    '分群於標準化空間進行,呈現時還原為原始單位(表四、圖二、圖六)。聚類為描述性技術,不主張因果。',

    '四群的整體特徵如下。低度參與群共 70 人、佔全體近三分之二,三項指標均為最低,使用率僅 0.102,是組織的多數常態。高投入待認可群共 25 人,使用率 0.674 達平均 2.2 倍,但認可度僅 0.36 且無人帶隊,是三項指標落差最大的一群。受認可的專業骨幹共 8 人,使用率與認可度均居次高,惟尚未擔任負責人。帶隊核心群共 4 人,三項指標同時最高,是唯一實際帶隊的群組。四群區隔並非由單一指標造成:低度參與群與其餘三群的差異在使用率,後三群的差異則依序落在認可度與領導角色。',

    '分受眾而言,三者處理的問題不同。人力資源部門面對級距與公平性:四群可對應四級獎金級距(圖七),辯護基礎在於級距由 KPI 實際分布劃分。部門主管面對人力配置:專業骨幹是負責人職位的候補來源,高投入待認可群則幾乎未獲指派。高階主管面對結構性風險:96% 從未帶隊。',

    '最後須反思。認可度與領導角色記錄的是組織給予的機會,非顧問單方面的產出;高投入待認可群認可度偏低,可能反映指派制度而非能力,連命名都有管理後果。本次分析亦修正了我對聚類方法的認識:原以為難處在演算法,實際上演算法僅數行,真正決定結果的是前處理與 k 的裁決——同一份資料在縮放與否之間,即得出兩種互不相容的分群。',

]

BODY['四、績效最優群組的識別'] = [
    '群 1「帶隊核心群」為績效最優群,共 4 人(員工編號 78、86、98、104)。',

    '依據為質心表的逐欄比較(表五):三項指標同時排名第一,使用率與認可度為全體平均的2.6 與 7.4 倍。領導角色不採倍數——其平均僅 0.056、96% 為 0,近零分母會誇大差異;改以絕對數:樣本 107 人中僅 4 人有帶隊紀錄,恰為此群,職位合計 6 個。惟領導角色本身即為分群變數,此屬描述性特徵,不構成獨立驗證。',

    '第三節指出該二指標帶有「機會」成分,故須確認判定不倚賴它們:僅以使用率排序,群 1 仍為最高 0.810(notebook §4.2b),三種權重方案亦皆最高,成員在 20 組種子下未變動。群 2 若政策不將帶隊列為必要條件,可與群 1 併級(圖七)。',

]

BODY['五、方法限制'] = [
    '序位變數被視為等距是固有限制:以秩轉換重跑 ARI 僅 0.38,群組邊界會移動,惟核心發現不變。資料亦僅單一時點。',

]

BODY['結論'] = [
    '採用方案為:排除員工識別碼、施以標準化、以 k = 4 分群。標準化由三項證據支持:認可度獨佔 79.6% 距離權重、分群退化為單變數複製、10 位顧問遭錯置;k = 4 經事前判準與敏感度檢驗。',

    '建議下一步:確認級距配置比例(圖七)、檢視指派流程、將專業骨幹納入接班規劃。',

]

REFS = [
    "Hubert, L., & Arabie, P. (1985). Comparing partitions. *Journal of Classification, 2*(1), "
    "193–218. https://doi.org/10.1007/BF01908075",

    "Milligan, G. W., & Cooper, M. C. (1988). A study of standardization of variables in cluster "
    "analysis. *Journal of Classification, 5*(2), 181–204. https://doi.org/10.1007/BF01897163",

    "Rousseeuw, P. J. (1987). Silhouettes: A graphical aid to the interpretation and validation of "
    "cluster analysis. *Journal of Computational and Applied Mathematics, 20*, 53–65. "
    "https://doi.org/10.1016/0377-0427(87)90125-7",
]

# ═══════════════════════════ 表格資料 ═══════════════════════════
T1 = (["變數", "業務意義", "型別", "範圍", "標準差", "零值佔比",
       "佔距離平方權重(未縮放)", "決策與理由"],
      [["EmployeeID", "員工識別碼(流水號)", "整數", "1–107", "30.89", "0%", "99%+",
        "排除:標稱型識別碼,數值無序位意義,且獨佔距離"],
       ["UsageRate", "在高優先級專案實際投入的時間佔比", "小數", "0.00–0.99", "0.3145", "10%", "10.7%",
        "納入:反映個人投入程度"],
       ["Recognition", "被特別指派參與的專案數量", "整數", "0–4", "0.8595", "80%", "79.6%",
        "納入:反映獲得指派的機會"],
       ["Leader", "擔任專案負責人的專案數量", "整數", "0–2", "0.3005", "96%", "9.7%",
        "納入:反映承擔的責任層級"]])

T2 = (["尺度", "該 10 人所屬群人數", "該群使用率跨幅", "業務可解釋性"],
      [["未縮放", "86", "0.00–0.92", "否:投入 92% 與 0% 同級距"],
       ["標準化 z 分數", "25", "0.41–0.94", "是:自成「高投入待認可群」"],
       ["MinMax 0–1", "25", "0.41–0.94", "是:與 z 分數結果完全相同"]])

T3 = (["k", "判準1 輪廓係數", "判準2 平方和降幅", "判準3 最小群人數",
       "判準4 群內認可度全距 ÷ 全體", "種子平均 ARI", "裁決"],
      [["2", "0.6435", "—", "24", "100%  ✗", "0.801", "否(判準4)"],
       ["3", "0.6703", "55.7%", "4", "100%  ✗", "1.000", "否(判準4)"],
       ["4", "0.6677", "44.0%", "4", "50%  ✓", "0.961", "採用"],
       ["5", "0.6823", "29.3%", "2  ✗", "50%  ✓", "0.980", "否(判準3)"],
       ["6", "0.6708", "30.2%", "2  ✗", "50%  ✓", "0.988", "否(判準3)"]])

T4 = (["群組", "人數", "佔比", "使用率", "認可度", "領導角色", "特徵"],
      [["低度參與群", "70", "65.4%", "0.102", "0.00", "0.0", "投入度為全體平均三分之一"],
       ["高投入待認可群", "25", "23.4%", "0.674", "0.36", "0.0", "投入高但幾乎未獲指派"],
       ["受認可的專業骨幹", "8", "7.5%", "0.751", "2.50", "0.0", "投入與認可皆高,尚未帶隊"],
       ["帶隊核心群", "4", "3.7%", "0.810", "2.75", "1.5", "三項全數領先"],
       ["(全體平均)", "107", "100%", "0.311", "0.374", "0.056", "—"]])

T5 = (["群組", "使用率(排名)", "認可度(排名)", "領導角色(排名)", "等權分數", "重領導", "重投入"],
      [["帶隊核心群", "0.810(1)", "2.75(1)", "1.50(1)　樣本內 6 個帶隊職位全數", "0.752", "0.745", "0.765"],
       ["受認可的專業骨幹", "0.751(2)", "2.50(2)", "0.00(3)", "0.461", "0.339", "0.567"],
       ["高投入待認可群", "0.674(3)", "0.36(3)", "0.00(3)", "0.257", "0.163", "0.367"],
       ["低度參與群", "0.102(4)", "0.00(4)", "0.00(3)", "0.034", "0.021", "0.052"]])

TOC_ENTRIES = ["管理摘要", "一、主要數據特徵:納入與排除的決策", "二、最優聚類數的決定",
               "三、四個群組的語言解釋", "四、績效最優群組的識別", "五、方法限制",
               "結論", "參考文獻"]


# ═══════════════════════════ 組裝 ═══════════════════════════
def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), EAST)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
    add_page_footer(doc)

    # ── 封面 ──
    for _ in range(4):
        para(doc, "")
    para(doc, "以 K-means 非監督式學習分析顧問績效", 20, True, WD_ALIGN_PARAGRAPH.CENTER, 4)
    para(doc, "—— IBM 獎金級距設計建議", 14, False, WD_ALIGN_PARAGRAPH.CENTER, 30)
    for line in ["評估任務 2b", "課程:421104 企業人工智慧",
                 "姓名:Po-Kai Huang    學號:26254793",
                 "繳交日期:2026 年 8 月 2 日"]:
        para(doc, line, 12, False, WD_ALIGN_PARAGRAPH.CENTER, 6)
    para(doc, "")
    para(doc, "附件:%s" % NB, 10.5, False, WD_ALIGN_PARAGRAPH.CENTER, 6)
    doc.add_page_break()

    # ── 目錄 ──
    para(doc, "目錄", 14, True, space_after=8)
    add_toc(doc, TOC_ENTRIES)
    doc.add_page_break()

    # ── 管理摘要 ──
    heading(doc, "管理摘要")
    for t in BODY["管理摘要"]:
        para(doc, t)

    fig(doc, "圖五　分析流程與各階段的關鍵判斷",
        "figures/fig5_workflow.png",
        "註:取自 notebook §3.6a。各節點標示的數值均為本次分析的實際結果,"
        "非流程示意。全程固定 random_state = 42、n_init = 10。")

    # ── 第一節 ──
    heading(doc, "一、主要數據特徵:納入與排除的決策")
    for t in BODY["一、主要數據特徵:納入與排除的決策"]:
        para(doc, t)
    table(doc, "表一　變數審核與納入決策",
          T1[0], T1[1],
          "註:數值取自 notebook §1.2 逐欄審核表與 §1.5 權重計算;業務意義依作業說明所定義。"
          "「佔距離平方權重」指該欄變異數佔歐氏距離平方總變異的比例。"
          "保留的三項指標分別涵蓋投入、機會與責任三個構面,共同構成績效的完整面貌。")
    table(doc, "表二　縮放決策對 10 位高投入顧問的實質影響",
          T2[0], T2[1],
          "註:取自 notebook §1.7。該 10 人定義為使用率 ≥ 0.60 且認可度 = 0"
          "(員工編號 2、44、47、60、77、81、88、91、96、97)。k 固定為 4。")
    fig(doc, "圖三　是否縮放,決定了 10 位高投入顧問被歸入哪一群",
        "figures/fig3_scaling_effect.png",
        "註:取自 notebook §3.5。左圖未縮放,10 位高投入者(紅色區域)被併入低投入群;"
        "右圖標準化後自成一群。認可度為整數,已加隨機抖動以避免點重疊。")

    # ── 第二節 ──
    heading(doc, "二、最優聚類數的決定")
    for t in BODY["二、最優聚類數的決定"]:
        para(doc, t)
    table(doc, "表三　四項判準的掃描結果與裁決",
          T3[0], T3[1],
          "註:取自 notebook §2.1、§2.2、§2.4、§2.5。四項判準於執行掃描前訂定,未因結果調整門檻。"
          "判準4 為群內單一 KPI 全距佔全體全距的比例,達 100% 者表示該指標在群內無區分力。"
          "ARI 以 20 組隨機種子(100–119)與 random_state = 42 的結果比對。k = 7 至 9 同樣含 2 人群,已略。"
          "判準3 的人數門檻無外部依據,已於 notebook §2.5c 做敏感度檢驗:取 3 或 4 人結論同為 k = 4。")
    fig(doc, "圖一　輪廓係數與肘部法對 k 的建議不一致",
        "figures/fig1_k_selection.png",
        "註:取自 notebook §2.3。輪廓係數(藍,左軸)最高點在 k = 5,"
        "群內平方和(紅,右軸)的轉折在 k = 4;兩法不一致時以業務可操作性裁決。")

    # ── 第三節 ──
    heading(doc, "三、四個群組的語言解釋")
    for t in BODY["三、四個群組的語言解釋"]:
        para(doc, t)
    table(doc, "表四　質心表(原始單位)與群組畫像",
          T4[0], T4[1],
          "註:取自 notebook §3.1、§3.3。分群於標準化空間進行,"
          "本表將各群還原為原始欄位的平均值以便解讀;標準化空間的質心 z 分數見 notebook §3.1 對照表。")
    fig(doc, "圖二　四個群組的績效輪廓(0 = 全體平均)",
        "figures/fig2_cluster_profile.png",
        "註:取自 notebook §3.4。縱軸單位為標準差;"
        "「高投入待認可群」的使用率明顯高於平均,但認可度與領導角色皆在平均之下。")

    fig(doc, "圖六　四個群組的畫像",
        "figures/fig6_cluster_cards.png",
        "註:取自 notebook §3.6b。四張卡共用 0–100% 同一刻度,條長為該群質心佔該指標全體最大值的百分比,"
        "黑色短線為全體平均,條末數字為原始單位的質心值。"
        "此處不採「相對全體平均的倍數」——領導角色全體平均僅 0.056,倍數會達 26.8 而使四張卡刻度不一致、無法比較。")

    # ── 第四節 ──
    heading(doc, "四、績效最優群組的識別")
    for t in BODY["四、績效最優群組的識別"]:
        para(doc, t)
    table(doc, "表五　質心逐欄比較與權重敏感度檢查",
          T5[0], T5[1],
          "註:取自 notebook §4.1、§4.3。排名 1 為該指標最高。"
          "三種權重方案的分數係將質心以 min-max 轉為 0–1 後加權求和;三方案結論一致,均指向帶隊核心群。"
          "領導角色不列倍數:其全體平均僅 0.056(96% 為 0),以近零數值為分母會誇大差異,故改列絕對數。"
          "另註:領導角色本身即為分群變數,該群集中持有帶隊職位屬描述性特徵,不構成獨立驗證。")
    fig(doc, "圖七　四級獎金級距、配置建議與對應行動",
        "figures/fig7_bonus_tiers.png",
        "註:取自 notebook §3.6c。人數為 k = 4 分群實算結果;"
        "配置佔比與行動為依分析結論提出的建議值,非模型輸出,須由人力資源部門依獎金政策核定。"
        "配置刻意不與人數成正比——最高級距僅 4 人,按人頭配置將使涵蓋面過窄。")

    fig(doc, "圖四　各群個體輪廓值分布(k = 4)",
        "figures/fig4_silhouette.png",
        "註:取自 notebook §5.4。107 人中僅員工 71 的輪廓值為負(−0.103),"
        "落在「低度參與群」與「高投入待認可群」的邊界,與 §2 的種子穩定性分析一致;"
        "紅色虛線為整體平均 0.668。")

    # ── 第五節、結論 ──
    heading(doc, "五、方法限制")
    for t in BODY["五、方法限制"]:
        para(doc, t)

    heading(doc, "結論")
    for t in BODY["結論"]:
        para(doc, t)

    # ── 參考文獻 ──
    heading(doc, "參考文獻")
    for r in REFS:
        p = doc.add_paragraph()
        set_font(p.add_run(r.replace("*", "")), 10.5)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.space_after = Pt(8)

    doc.save(OUT)
    return doc


if __name__ == "__main__":
    build()
    body_texts = [t for k, v in BODY.items() for t in v]
    headings = list(BODY.keys()) + ["參考文獻"]
    strict = count_units(body_texts + headings)
    loose = strict + count_units(NOTES)
    print("已產出:%s" % os.path.basename(OUT))
    print("字數(嚴格:內文 + 標題,不含表圖註/參考文獻/封面):%d" % strict)
    print("字數(寬鬆:再加表圖註):%d" % loose)
    print("目標區間:1350–1650(1500 ±10%)")
    print("嚴格口徑落點:%s" % ("在區間內 OK" if 1350 <= strict <= 1650 else "!! 超出區間"))
