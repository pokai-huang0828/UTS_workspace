# -*- coding: utf-8 -*-
# A2a 報告產生器 v3
#   v2 = Codex round1 十八條缺陷全修
#   v3(2026-07-25)= 引用誤用修正 + 量化成本表 + FP 商業意義 + recall 論證
#                    + 正/負向用詞 + 統計措辭修正 + 圖表註壓縮
#   依據:Codex 對抗審查(引用查證 3/3 成立)、7/22 管濟偉 Zoom、7/23 李春平 Zoom、A1 復盤評語①
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EAST = '微軟正黑體'

def set_font(run, size=11, bold=False, italic=False):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), EAST)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def para(doc, text, size=11, bold=False, align=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size, bold)
    if align: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

def heading(doc, text):
    h = doc.add_heading(level=1)
    r = h.add_run(text)
    set_font(r, 14, True)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); f1.set(qn('w:dirty'), 'true')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'TOC \\o "1-1" \\h \\z \\u'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'separate')
    # 欄位快取寫入實際條目:Word 開啟仍自動更新,Canvas 預覽器則顯示靜態目錄(雙保險)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve')
    t.text = ('管理摘要　|　一、分類模型的評估與推薦　|　二、最佳模型誤判造成的收入損失　|　'
              '三、最重要屬性的識別與領域解釋　|　四、降低客戶流失率的策略　|　結論　|　參考文獻')
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for el in (f1, it, f2, t, f3): r._r.append(el)

def add_page_footer(doc):
    sec = doc.sections[0]
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    f3 = OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'), 'end')
    for el in (f1, it, f3): r._r.append(el)
    set_font(r, 9)

NOTES = []          # 收集所有圖表註,供字數雙口徑統計

def table(doc, caption, headers, rows, source_note, font_size=9.5):
    NOTES.append(source_note)
    cap = doc.add_paragraph(); rc = cap.add_run(caption); set_font(rc, 10, True)
    cap.paragraph_format.space_after = Pt(3)
    cap.paragraph_format.keep_with_next = True
    tb = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tb.style = 'Table Grid'
    trPr = tb.rows[0]._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)
    for row in tb.rows:
        rp = row._tr.get_or_add_trPr()
        cs = OxmlElement('w:cantSplit'); rp.append(cs)
    for j, h in enumerate(headers):
        r = tb.rows[0].cells[j].paragraphs[0].add_run(h); set_font(r, font_size, True)
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            r = tb.rows[i].cells[j].paragraphs[0].add_run(str(v)); set_font(r, font_size)
    note = doc.add_paragraph(); r = note.add_run(source_note); set_font(r, 8.5, italic=True)
    note.paragraph_format.space_after = Pt(10)
    return tb

def fig(doc, caption, path, source_note):
    NOTES.append(source_note)
    cap = doc.add_paragraph(); rc = cap.add_run(caption); set_font(rc, 10, True)
    cap.paragraph_format.space_after = Pt(3)
    cap.paragraph_format.keep_with_next = True
    doc.add_picture(path, width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.keep_with_next = True
    note = doc.add_paragraph(); r = note.add_run(source_note); set_font(r, 8.5, italic=True)
    note.paragraph_format.space_after = Pt(10)

def count_units(texts):
    # Word 口徑近似:中日韓字元(含全形標點)逐字計 1,英文/數字連續串計 1
    n = 0
    for t in texts:
        n += len(re.findall('[一-鿿　-〿＀-￯]', t))
        n += len(re.findall('[A-Za-z0-9]+', t))
    return n

# ─────────────────────────── 內文 ───────────────────────────
BODY = {}

BODY['管理摘要'] = (
"本報告為電信公司建立客戶流失(Churn)預測的概念驗證。以 3,333 筆數據比較六種分類模型,"
"依事前訂定的規則推薦 XGBoost(測試集 accuracy 0.925、F1 0.715;RandomForest 為次選)。"
"該模型漏判 51 名流失客戶(False Negative),十二個月毛收入暴露上限約 A$35,926。"
"最重要變數為日間通話分鐘數(正向),月費與客服來電次數次之(正向),合約續約為負向;"
"第四節據此提出三項行動與量化商業案例。")

BODY['一、分類模型的評估與推薦'] = [
("方法設計:分層切分(70/30)、以 Pipeline 使縮放只在訓練折內擬合、選模僅用訓練集的 25 折重複交叉驗證,"
"測試集不參與選模(見附件 notebook)。全猜「不流失」也有 85.5% 的 accuracy,accuracy 因此具誤導性。"
"選模主指標為 F1,因它同時約束漏判與誤報:單看 recall 會選到 SVC(表一,recall 0.807、precision 僅 0.498),"
"其誤報負擔明顯較高,是否可承受須依單次接觸成本與客服產能評估。"
"惟就商業風險而言 recall 最關鍵——流失者僅佔 14.5%,漏掉一位即一筆無人挽留的流失,"
"故 recall 亦為事前訂定的平手裁決指標;precision 則用於監控挽留資源浪費。"),
("結果如表一與圖一:XGBoost 與 RandomForest 領先其餘四模型,交叉驗證 F1 幾乎相同(0.732 對 0.732);"
"經重複交叉驗證的校正檢定(Nadeau & Bengio, 2003),三項指標皆未偵測到顯著差異——"
"未偵測到差異不等於證明等效,僅代表本樣本無法區分。故依事前訂定的平手規則,"
"取 recall 均值較高的 XGBoost(0.656 對 0.639)為最終模型,此為規則選定而非性能勝出。"
"測試集上 RandomForest 的 F1(0.745)較高,惟成對 bootstrap 95% 區間 [−0.001, +0.064] 與零相容,依協議不回選。"
"建模方法造成的準確率差異可使挽留活動獲利相差甚鉅(Neslin et al., 2006),故選模與最終評估分離。")]

BODY['二、最佳模型誤判造成的收入損失'] = [
("混淆矩陣(圖二)的兩類誤判成本方向相反。False Negative(FN)是實際流失卻被預測為留存:"
"若模型標記是挽留的觸發條件,這些客戶不會被接觸,形成未挽留的收入暴露。"
"False Positive(FP)相反——實際留存卻被判為流失,公司投入不必要的工時與優惠;XGBoost 為 24 人,"
"其可承受程度需結合單次接觸成本與客服產能評估。本案以 F1 平衡兩類錯誤,並以 recall 監控 FN 風險。"),
("XGBoost 於測試集(1,000 筆,含 145 名流失者)產生 51 名 FN,實際月費合計 A$2,993.8(平均每人每月 A$58.7)。"
"表二列三種留存假設:十二個月基準下毛收入暴露上限約 A$35,926;"
"外推至全體流失者,預估漏抓約 170 人、年暴露約 A$119,669,屬毛收入口徑的量級參考;淨效益見第四節。")]

BODY['三、最重要屬性的識別與領域解釋'] = [
("各模型原生重要度排序不一致,故以最終模型在測試集上的 permutation importance 裁決(重複 30 次、F1 評分):"
"DayMins 30 次全數第一;MonthlyCharge 與 CustServCalls 互有領先(23/30),屬同層級第二梯隊;"
"ContractRenewal 多數第四(圖三)。此一致性限於同一已擬合模型與測試集,非跨樣本穩健性宣稱。"),
("影響方向可由群組平均差異讀出(表三):DayMins、MonthlyCharge、CustServCalls 為正向"
"(流失群組平均較高);ContractRenewal 為負向(已續約者流失明顯較低)。"
"此為兩群的描述性差異,非控制其他變數後的效果,亦不宣稱單調或因果關係。"),
("領域解釋(與資料一致的假說,非因果結論):高用量客戶帳單高,可能對資費與體驗較敏感;"
"同型電信資料研究亦將 total day minutes、total day charge 與 customer service calls 列為前三特徵"
"(Abdelhady & Mohamed, 2025)。頻繁致電客服是不滿信號;未續約者轉換成本低。"
"DayMins 與 MonthlyCharge 中度相關(r = 0.57),宜視為同一「高用量-高帳單」信號。")]

BODY['四、降低客戶流失率的策略'] = [
("挽留策略完全建立於第三節識別的變數:每月以 XGBoost 評分產生高風險名單,"
"再依第三節變數建立三組試點候選標籤(表四)。三組以各自門檻獨立定義、可能重疊,"
"表五僅呈現各行動單獨實施的情境,不得加總;若同時試行,重疊客戶依一、二、三順序只分派一次。"
"門檻用於招募試點對象,不代表模型分數本身。"),
("行動一(CustServCalls)服務恢復:來電達門檻即觸發專員回訪與問題升級,處理可能的服務問題。"
"行動二(ContractRenewal)續約轉化:向未續約的高風險客戶提供限時續約優惠。"
"行動三(DayMins/MonthlyCharge)資費重組:主動推薦封頂或客製方案。"),
("表五為情境試算(假設毛利率 60%、每名挽回客戶保留十二個月毛利 A$423):三項行動的損益兩平增量留存率"
"分別為 13.6%、3.2%、9.1%,於 20% 增量留存假設下皆為正淨效益。"
"純以門檻排序應為二、三、一;若客服產能允許,建議行動一先行(無誘因成本)、行動二緊隨。"
"各行動均以隨機對照組試點,三個月檢視預設 KPI,未優於對照即停止。")]

BODY['結論'] = (
"本概念驗證顯示:於本次切分,現有變數已可支持具辨識能力的模型,XGBoost 為依事前規則選定者,"
"跨期穩定性仍須以時間外樣本驗證。FN 代表情境化的收入暴露,而非已證實可挽回的損失;"
"三項行動的損益兩平門檻見表五,可達性須由隨機對照試點估計。"
"若客服產能允許,建議先以行動一進行小型對照試點,KPI 達標後才擴大。")

# ─────────────────────────── 組裝 ───────────────────────────
doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn('w:eastAsia'), EAST)
add_page_footer(doc)

# 封面
for _ in range(5): doc.add_paragraph()
para(doc, '電信客戶流失預測概念驗證報告', 20, True, WD_ALIGN_PARAGRAPH.CENTER)
para(doc, '評估任務二a:使用監督學習技術分析數據', 14, False, WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
for line in ('課程:421104 Artificial Intelligence for Enterprises(企業人工智能)',
             '姓名:黃柏凱(PoKai Huang)', '學號:26254793',
             '參考格式:APA 第 7 版', '日期:2026 年 7 月',
             '附件:Huang_26254793_421104_Assessment 2a.ipynb(完整程式與輸出)'):
    para(doc, line, 12, False, WD_ALIGN_PARAGRAPH.CENTER, 4)
doc.add_page_break()

# 目錄
para(doc, '目錄', 14, True)
add_toc(doc)
doc.add_page_break()

# 管理摘要
heading(doc, '管理摘要')
para(doc, BODY['管理摘要'])

# 一、
heading(doc, '一、分類模型的評估與推薦')
para(doc, BODY['一、分類模型的評估與推薦'][0])
table(doc, '表一、六模型於 5×5 重複分層交叉驗證與保留測試集之表現(churn 類指標)',
    ['模型', 'CV F1(mean±std)', '測試 Accuracy', '測試 Precision', '測試 Recall', '測試 F1'],
    [['XGBoost(推薦)', '0.732 ± 0.053', '0.925', '0.797', '0.648', '0.715'],
     ['RandomForest(次選)', '0.732 ± 0.053', '0.935', '0.864', '0.655', '0.745'],
     ['SVC(balanced)', '0.652 ± 0.034', '0.854', '0.498', '0.807', '0.616'],
     ['DecisionTree', '0.625 ± 0.047', '0.878', '0.578', '0.586', '0.582'],
     ['KNN', '0.555 ± 0.042', '0.892', '0.713', '0.428', '0.535'],
     ['LogisticRegression(balanced)', '0.483 ± 0.028', '0.765', '0.356', '0.766', '0.486']],
    '註:notebook §8–9。CV 僅用訓練集。')
para(doc, BODY['一、分類模型的評估與推薦'][1])
fig(doc, '圖一、六分類器於保留測試集之表現', 'figures/fig1_model_comparison.png',
    '註:notebook §9。虛線為全猜不流失基準。')

# 二、
heading(doc, '二、最佳模型誤判造成的收入損失')
para(doc, BODY['二、最佳模型誤判造成的收入損失'][0])
fig(doc, '圖二、XGBoost 測試集混淆矩陣(FN = 51)', 'figures/fig2_confusion_winner.png',
    '註:notebook §10。')
para(doc, BODY['二、最佳模型誤判造成的收入損失'][1])
table(doc, '表二、FN 收入暴露三層估算(依 FN 個體實際月費)',
    ['留存假設', '測試集 FN 毛暴露上限', '全資料集預估 FN 暴露(推估)'],
    [['6 個月', 'A$17,963', 'A$59,835'],
     ['12 個月(基準)', 'A$35,926', 'A$119,669'],
     ['24 個月', 'A$71,851', 'A$239,339']],
    '註:notebook §11。毛收入口徑;全資料集欄為外推推估。')

# 三、
heading(doc, '三、最重要屬性的識別與領域解釋')
para(doc, BODY['三、最重要屬性的識別與領域解釋'][0])
fig(doc, '圖三、Permutation importance(測試集,重複 30 次,scoring=F1)', 'figures/fig3_perm_importance.png',
    '註:notebook §13。橫軸為 F1 平均降幅(±SD)。')
para(doc, BODY['三、最重要屬性的識別與領域解釋'][1])
table(doc, '表三、重要變數於流失/未流失群組的平均值差異與影響方向',
    ['變數', '未流失(Churn=0)', '流失(Churn=1)', '影響方向'],
    [['DayMins(日間通話分鐘)', '175.2', '206.9', '正向(越高越易流失)'],
     ['MonthlyCharge(月費, A$)', '55.8', '59.2', '正向(越高越易流失)'],
     ['CustServCalls(客服來電次數)', '1.45', '2.23', '正向(越多越易流失)'],
     ['ContractRenewal(續約率)', '93.5%', '71.6%', '負向(已續約者較不流失)']],
    '註:notebook §14。全體群組平均。')
para(doc, BODY['三、最重要屬性的識別與領域解釋'][2])

# 四、
heading(doc, '四、降低客戶流失率的策略')
para(doc, BODY['四、降低客戶流失率的策略'][0])
table(doc, '表四、三項行動之試點設計',
    ['行動(對應變數)', '試點門檻', '覆蓋客群', '該群流失率(倍數)', '負責單位', '上線順序'],
    [['一、服務恢復\n(CustServCalls)', '來電 ≥3 次', '696 人(21%)', '26.1%(1.8×)', '客服部', '1'],
     ['二、續約轉化\n(ContractRenewal)', '未續約', '323 人(10%)', '42.4%(2.9×)', '行銷部', '2'],
     ['三、資費重組\n(DayMins/MonthlyCharge)', '日通話 ≥216 分且月費 ≥ 中位(A$53.5)', '767 人(23%)', '31.3%(2.2×)', '資費部', '3']],
    '註:notebook 實算。倍數 = 該群流失率 ÷ 全體 14.5%。', font_size=9)
para(doc, BODY['四、降低客戶流失率的策略'][1])
table(doc, '表五、三項行動之量化商業案例(損益兩平與淨效益試算)',
    ['行動', '每人接觸成本', '成功挽留誘因成本', '試點總投入', '該群預期流失人數',
     '損益兩平所需增量留存率', '淨效益 @20% 增量留存'],
    [['一、服務恢復', 'A$15', '—', 'A$10,440', '182 人', '13.6%', '+A$4,957'],
     ['二、續約轉化', 'A$5', 'A$59(一個月折抵)', 'A$1,615', '137 人', '3.2%', '+A$8,359'],
     ['三、資費重組', 'A$10', 'A$70(月費調降 10%×12)', 'A$7,670', '240 人', '9.1%', '+A$9,274']],
    '註:notebook 補充二。成本為估計值,誘因僅於挽留成功時發生。', font_size=8.5)
para(doc, BODY['四、降低客戶流失率的策略'][2])

# 結論
heading(doc, '結論')
para(doc, BODY['結論'])

# 參考文獻(期刊名+卷斜體)
heading(doc, '參考文獻')
REFS = [
 [('Abdelhady, M. G., & Mohamed, K. A. (2025). Leveraging artificial intelligence for predictive customer churn modeling in telecommunications: A framework for enhanced customer relationship management. ', False),
  ('Scientific Reports, 15', True), (', Article 43826. https://doi.org/10.1038/s41598-025-30108-z', False)],
 [('Nadeau, C., & Bengio, Y. (2003). Inference for the generalization error. ', False),
  ('Machine Learning, 52', True), ('(3), 239–281. https://doi.org/10.1023/A:1024068626366', False)],
 [('Neslin, S. A., Gupta, S., Kamakura, W., Lu, J., & Mason, C. H. (2006). Defection detection: Measuring and understanding the predictive accuracy of customer churn models. ', False),
  ('Journal of Marketing Research, 43', True), ('(2), 204–211. https://doi.org/10.1509/jmkr.43.2.204', False)],
]
for segs in REFS:
    p = doc.add_paragraph()
    for text, ital in segs:
        r = p.add_run(text); set_font(r, 10.5, italic=ital)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after = Pt(6)

OUT = 'Huang_26254793_421104_Assessment 2a.docx'
doc.save(OUT)

# ── 字數雙口徑 ──
HEADS = ['管理摘要', '一、分類模型的評估與推薦', '二、最佳模型誤判造成的收入損失',
         '三、最重要屬性的識別與領域解釋', '四、降低客戶流失率的策略', '結論']
strict = count_units(HEADS)
for k in ('管理摘要', '結論'):
    strict += count_units([BODY[k]])
for k in HEADS[1:5]:
    strict += count_units(BODY[k])
notes_n = count_units(NOTES)
print(f'SAVED: {OUT}')
print(f'嚴格口徑(正文+標題,不含表格/圖說/註/文獻):{strict}')
print(f'圖表註合計:{notes_n}')
print(f'寬鬆口徑(正文+標題+註):{strict + notes_n}')
print(f'規定 1350–1650 —— 兩種口徑都必須落在區間內')
