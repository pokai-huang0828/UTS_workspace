# -*- coding: utf-8 -*-
"""Build the A3 Word report (26254793_A3.docx) with required formatting + embedded figures.
Re-run to regenerate as sections are added. Body text lives here = single source of truth."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

BASE = Path(__file__).resolve().parent
FIG = BASE / "figures"
FONT = "Microsoft JhengHei"

doc = Document()

# ---- page + base style ----
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)            # A4
sec.top_margin = sec.bottom_margin = Cm(2.5)
sec.left_margin = sec.right_margin = Cm(2.5)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)
_rpr = normal.element.get_or_add_rPr()
_rf = _rpr.find(qn("w:rFonts"))
if _rf is None:
    _rf = OxmlElement("w:rFonts"); _rpr.append(_rf)
for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
    _rf.set(qn(a), FONT)

def _font(run, size, bold=False, italic=False):
    run.font.name = run.font.name or FONT
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), FONT)

CHARS = []
def title(t, size=16):
    p = doc.add_paragraph(); p.alignment = AL.CENTER
    _font(p.add_run(t), size, bold=True)
def heading(t, size=14):
    p = doc.add_paragraph(); p.alignment = AL.LEFT; p.paragraph_format.space_before = Pt(10)
    _font(p.add_run(t), size, bold=True)
def body(t):
    CHARS.append(t)
    p = doc.add_paragraph(); p.alignment = AL.JUSTIFY
    _font(p.add_run(t), 12)
def figure(fname, caption, width=14.5):
    p = doc.add_paragraph(); p.alignment = AL.CENTER; p.paragraph_format.space_before = Pt(4)
    p.add_run().add_picture(str(FIG / fname), width=Cm(width))
    c = doc.add_paragraph(); c.alignment = AL.CENTER; c.paragraph_format.space_after = Pt(10)
    _font(c.add_run(caption), 10, italic=True)

# ================= CONTENT =================
def center(text, size, bold=False, italic=False):
    p = doc.add_paragraph(); p.alignment = AL.CENTER
    _font(p.add_run(text), size, bold=bold, italic=italic)

# ---- cover page ----
for _ in range(5):
    doc.add_paragraph()
center("UTS 231708 — Foundation Studio(Session 3, 2026)", 12)
doc.add_paragraph()
center("評估三:商業分析項目", 22, bold=True)
center("NIB Holdings 案例分析", 16, bold=True)
for _ in range(6):
    doc.add_paragraph()
center("PoKai Huang(黃柏凱)", 14)
center("學號:26254793", 12)
doc.add_page_break()

heading("引言")
body("這份報告把我放在 NIB 一位資料分析師的位置。公司的執行長有個擔憂:隨著醫療成本逐年墊高,加上政府未來有可能縮減鼓勵民眾投保的稅務誘因,NIB 恐怕得被迫調漲保費。她想先弄清楚一件事——如果保費往上調 10%,需求會掉多少?")
body("本報告分三步回答。先盤點 NIB 在整個私人健康保險市場的位置與表現(第一部分);接著回顧實證文獻,看價格變動對保險需求到底有多大影響(第二部分);最後提出一套自己的評估方案,說明在只有觀測資料的情況下,該怎麼可信地估出漲價對需求的因果效應(第三部分)。")

heading("第一部分:公司分析")
body("澳洲的醫療體系是公私並行的。政府的 Medicare 替大多數人負擔了基本門診和必要的住院費用,但如果想要更短的候診時間、自己挑醫生、甚至住進單人房,就得另外買一份私人健康保險。截至 2026 年 3 月,全澳大約有 45.8% 的人持有住院險(APRA, 2026)——不過這個數字是熬過一段下坡才爬回來的:覆蓋率在 2015 年中見頂(47.3%),接著一路滑到 2020 年疫情期間的 43.7% 谷底,近兩年才逐步回升(圖一)。換句話說,NIB 所處的這個市場,需求面本身就帶著一點逆風。")
figure("fig1_coverage_trend.png", "圖一  澳洲住院險覆蓋率,2007–2026(來源:APRA 季度統計)")
body("NIB(ASX:NHF)是一家營利性的私人健康保險公司,2007 年於澳洲證券交易所上市,以保費收入計是全澳第四大的保險商。若把所有保險公司攤開來比,NIB 其實是個中段班。以 2024–25 年度的保費收入計,市場由 Bupa(25.5%)和 Medibank(24.7%)兩家領跑,接著是 HCF(13.2%),NIB 以 9.0% 排第四(圖二;APRA, 2025)。它的規模大概只有龍頭的三分之一,離市場主導者還有不小距離。")
figure("fig4_market_share.png", "圖二  各保險商市佔率(以保費收入計),2024–25(來源:APRA)")
body("但市場顯然不是只看規模。把 NIB 的股價和整個大盤(ASX 200)都從 2011 年設為 100 來看,NIB 一路漲到約 520,是 15 年前的五倍多,同期大盤只翻了一倍(圖三)。投資人願意給它這麼高的評價,看的不是它有多大,而是它賺錢的方式——這就帶到 NIB 真正的強項:利潤率。")
figure("fig6_share_price.png", "圖三  NIB 股價 vs ASX 200,2011 年 = 100(來源:Yahoo Finance)")
body("用 APRA 的口徑來算(淨利率 = 1 −(理賠 + 營運費用)÷ 保費收入),NIB 在 2024–25 年的健康險淨利率是 7.7%,不只贏過 Medibank(7.2%)和 Bupa(5.5%),更是幾家大型保險商裡最高的(圖四)。對照之下,幾家非營利的互助型基金就吃力得多——HCF 的承保淨利率甚至是 −0.8%,等於本業在賠錢、得靠投資收益來補。這個落差不是巧合:NIB 和 Medibank 這類上市營利公司,本來就背著股東對報酬的期待,在成本控制和訂價上會更斤斤計較。具體看,NIB 的管理費用率(營運費用 ÷ 保費收入)大約只有 11%,在同業裡偏低,這正是它淨利率能領先的主因。")
figure("fig5_net_margin.png", "圖四  大型保險商健康險淨利率,2024–25(來源:APRA)")
body("不過,把時間拉長就看得到壓力。NIB 的承保營運利潤(UOP,公司自己用來跨年比較的指標)從 FY21 的 2.05 億澳元一路爬到 FY23 的 2.63 億見頂,接著連兩年回落到 FY25 的 2.39 億(圖五;nib Holdings, 2021)。原因是這兩年醫療理賠成長得比保費還快——而這正是執行長擔心、想靠漲保費補回來的那股壓力。即便如此,NIB 的底線還算穩:FY25 稅後淨利 1.986 億澳元,比前一年又多了 9.4%(nib Holdings, 2025)。")
figure("fig7_nib_uop.png", "圖五  NIB 承保營運利潤(UOP),FY21–FY25(來源:NIB 年報)")
body("整體看,NIB 是一家規模不大、但很會賺的營利型保險商。它的競爭策略不是跟兩大龍頭硬拚市佔,而是守住利潤率、再靠海外學生和紐西蘭等業務找成長——它的澳洲健保會員數在 FY25 還逆勢成長了 3.2%。但它和整個產業一樣,得面對覆蓋率停滯、年輕族群慢慢流失的長期逆風——這也就把問題帶到了核心:萬一真的漲保費,需求會怎麼反應?")

heading("第二部分:文獻綜述")
body("執行長想知道的,其實是一個經濟學概念:私保需求的「價格彈性」——保費漲了,人會少投保多少。這得看實證證據,不能憑直覺。")
body("最直接的一篇是 Buchmueller 等人(2021)。他們拿澳洲的醫療保險附加稅(MLS,對高收入卻沒買私保者課的稅)當天然實驗:2008 年政府大幅調高課稅門檻,等於讓一批人突然不必再為省稅而投保。配上 HILDA 這份多年追蹤的家戶資料,他們發現被課 MLS 會讓當年投保機率上升 2 到 3 個百分點,長期持續被課則逐年累積、十年後高達 13 個百分點,而且年輕人的反應特別大。這個估計要能解讀成因果,前提是門檻邊緣剛好被課和剛好沒被課的人在其他條件上夠相似、沒有刻意壓低申報所得來避稅;在這個假設下它的識別很乾淨,但也因此只說明了門檻附近那群高收入者,未必能推廣到全體。")
body("Cheng(2014)的角度不同。他用聯立方程式,把「投保」和「就醫」這兩個互相牽動的決定一起估,模擬政府若減少保費補貼(實質上的漲價)會如何。結果是公共部門反而省錢——少付的補貼大於民眾轉回公立醫院的成本。和 Buchmueller 的天然實驗相比,Cheng 沒有現成的政策衝擊可借,因果結論完全押在「那組聯立方程式正確描述了投保與就醫的真實關係」這個假設上——模型設定錯了,結論就跟著錯;但它的好處是能直接模擬不同幅度的政策變動。")
body("至於保費為何被推著漲,Doiron 與 Kettlewell(2018)用工具變數估出,有私保會讓住院機率高出 5 到 6 個百分點,主因是把原本上公立醫院的病人拉到私立——也就是保險本身誘發了更多、也更貴的就醫。")
body("綜合來看:私保需求短期內對價格並不太敏感(MLS、終身保險加成、政府補貼都把人綁著),但反應會隨時間放大,而且最敏感的就是年輕、健康的那一群。所以保費真漲 10%,短期流失的保單可能有限、NIB 的保費收入甚至可能不減反增;但先走的會不成比例地是年輕人,把風險池愈推愈老。")

heading("第三部分:評估方案")
body("私保需求最脆弱的環節是年輕人。40 歲以下投保人的占比從 2015 年起一路下滑(圖六),這群人正是保費一漲最可能退保的邊際客戶。所以要評估漲價的需求效應,聚焦年輕族群最有意義。")
figure("fig2_under40_share.png", "圖六  40 歲以下投保人占總投保人數的比例,2007–2026(來源:APRA)")
body("問題是,這個效應很難算得乾淨。最直覺的做法是把各家公司歷年的保費和投保人數兜起來跑迴歸,但那個係數幾乎一定不是因果:保險公司是看著理賠和需求在訂價,價格和數量其實互相決定(反向因果);所得、健康、年齡、政策同時牽動保費和需求,漏掉就會偏(遺漏變數);而且會投保的人本來就和不投保的不一樣(選擇偏誤)。只有觀測資料,要繞過這些問題,關鍵是找到一個跟需求無關的保費外生變動。")
body("我的方案是用差異中之差異(DiD),搭上 2019 年 4 月上路的年齡折扣——政府讓 18 到 29 歲的人投保可以打折,等於只對年輕人降價。處理組取 25–29 歲、對照組取年紀相近卻沒享到折扣的 30–34 歲,用 APRA 各年齡層的季度投保人數,比較折扣前後兩組的變化差;寫成迴歸,就是在虛擬變數上加一個「年輕×折扣後」的交互項,那個係數就是降價的因果效應。這套設計成立的前提是「平行趨勢」:沒有折扣的話,兩組的投保人數本來會沿著同樣的趨勢走。")
center("投保人數 = β₀ + β₁·年輕 + β₂·折扣後 + β₃·(年輕 × 折扣後) + 控制 + ε      (β₃ = 漲/降價的因果效應)", 11, italic=True)
body("我實際拉資料跑了一遍。樸素的 DiD 算出來是 −4.1%,也就是折扣後年輕組反而掉得比對照組更多。但這個數字不能當因果——因為平行趨勢根本不成立:在政策上路前,25–29 歲的投保人數每季就已經以 1.5% 的速度下滑,是 30–34 歲(0.7%)的兩倍(圖七)。換句話說,年輕人本來就在加速流失,2019 那個最多只有 2% 的小折扣,完全沒能把這條下滑的線拉回來。這也正好提醒:沒檢查平行趨勢就把 DiD 的數字當因果,會得到完全相反的結論。")
figure("fig3_did_setup.png", "圖七  DiD 平行趨勢檢查:25–29 歲 vs 30–34 歲投保人數,2019 Q1 = 100(來源:APRA)")
body("正因為如此,這套方案把「先檢驗平行趨勢」當成第一道關卡——一旦像上面那樣不成立,就升級成事件研究法(event study),逐季估出兩組差異、把既有趨勢扣掉,或改挑趨勢更貼近的對照組。方案的優點是可行(資料現成)、識別邏輯清楚,又能誠實暴露限制;其他限制包括:2020 年的疫情和 2019 年同時上路的保單分級改革都會干擾;APRA 的資料是總量、也沒有個人所得,沒辦法直接套到 MLS 的收入門檻上。若要更進一步,可以改用 MLS 門檻做斷點回歸,或拿各家公司不同的調漲幅度當工具變數。")

heading("結論")
body("整體而言,NIB 是一家體質強健的營利型保險商,靠利潤率和多元成長在兩大龍頭旁邊站穩,但和整個產業一樣,面對覆蓋率停滯、人口老化的逆風。文獻顯示私保需求短期相對缺乏彈性、年輕人較敏感;而要可信地估出漲價對需求的影響,得靠準實驗設計,而不是單純的相關。對執行長的 10% 問題,我的判斷是:短期保費收入大概還守得住,但年輕客群加速流失、風險池惡化,才是真正該擔心的長期威脅。")

# ---- references (APA 7; clickable links; not counted toward word limit) ----
def add_hyperlink(paragraph, url, text=None):
    text = text or url
    r_id = paragraph.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(a), FONT)
    rpr.append(rf)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "24"); rpr.append(sz)
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1"); rpr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t); link.append(run); paragraph._p.append(link)

def ref(text, url=None):
    p = doc.add_paragraph(); p.alignment = AL.LEFT
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(4)
    _font(p.add_run(text + (" " if url else "")), 12)
    if url:
        add_hyperlink(p, url)

heading("參考文獻")
ref("Australian Prudential Regulation Authority. (2025). Annual private health insurance performance statistics 2024–25. APRA.",
    "https://www.apra.gov.au/operations-of-private-health-insurers-annual-report")
ref("Australian Prudential Regulation Authority. (2026). Quarterly private health insurance statistics, March 2026. APRA.",
    "https://www.apra.gov.au/quarterly-private-health-insurance-statistics")
ref("Buchmueller, T. C., Cheng, T. C., Pham, N. T., & Staub, K. E. (2021). The effect of income-based mandates on the demand for private hospital insurance and its dynamics. Journal of Health Economics, 75, 102403.",
    "https://doi.org/10.1016/j.jhealeco.2020.102403")
ref("Cheng, T. C. (2014). Measuring the effects of reducing subsidies for private insurance on public expenditure for health care. Journal of Health Economics, 33(1), 159–179.",
    "https://doi.org/10.1016/j.jhealeco.2013.11.007")
ref("Doiron, D., & Kettlewell, N. (2018). The effect of health insurance on the substitution between public and private hospital care. The Economic Record, 94(305), 135–154.",
    "https://doi.org/10.1111/1475-4932.12394")
ref("nib Holdings. (2021). Annual report 2021. nib Holdings Limited.",
    "https://www.nib.com.au/shareholders/annual-reports")
ref("nib Holdings. (2025). FY25 full year results [ASX announcement]. nib Holdings Limited.",
    "https://www.nib.com.au/media/company/fy25-asx-announcement")
ref("Yahoo Finance. (2026). nib Holdings Limited (NHF.AX) and S&P/ASX 200 (^AXJO) historical prices [Data set].",
    "https://finance.yahoo.com")

out = BASE / "26254793_A3.docx"
try:
    doc.save(str(out)); saved = out.name
except PermissionError:
    alt = BASE / "26254793_A3_preview.docx"
    doc.save(str(alt)); saved = alt.name
    print("(主檔在 Word 開著、被鎖 → 改存預覽檔)")
import re
_s = "".join(CHARS)
_cjk = len(re.findall(r"[一-鿿]", _s))
_tok = len(re.findall(r"[A-Za-z]+|\d+(?:\.\d+)?", _s))
print(f"Saved {saved}")
print(f"  raw chars: {len(_s)}  |  CJK 中文字: {_cjk}  |  Word-style(中文字+英數詞): ~{_cjk + _tok}  (target ~2500)")
