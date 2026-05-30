# -*- coding: utf-8 -*-
import copy
from docx import Document
from docx.oxml.ns import qn

SRC = "評估二_商業分析報告.docx"
doc = Document(SRC)
paras = doc.paragraphs

def last_run(p):
    return p.runs[-1] if p.runs else None

def append_text(p, text):
    """Append text as a new run cloning the formatting of the paragraph's last run."""
    base = last_run(p)
    new_r = copy.deepcopy(base._element)
    # set text
    for t in new_r.findall(qn('w:t')):
        new_r.remove(t)
    # remove any breaks/tabs copied
    for tag in ('w:br','w:tab','w:cr'):
        for el in new_r.findall(qn(tag)):
            new_r.remove(el)
    t = new_r.makeelement(qn('w:t'), {qn('xml:space'):'preserve'})
    t.text = text
    new_r.append(t)
    p._p.append(new_r)

def set_paragraph_text(p, text):
    """Replace all text in paragraph, keep first run's formatting."""
    runs = p.runs
    if not runs:
        return
    first = runs[0]
    # remove extra runs
    for r in runs[1:]:
        r._element.getparent().remove(r._element)
    first.text = text

def insert_para_before(ref_para, text, template_para):
    """Insert a new body paragraph (cloned style from template_para) before ref_para."""
    new_p = copy.deepcopy(template_para._p)
    # clear runs in clone
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    ref_para._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    np = Paragraph(new_p, ref_para._parent)
    # build a run cloning template's first run formatting
    tmpl_run = template_para.runs[0]._element if template_para.runs else None
    if tmpl_run is not None:
        r = copy.deepcopy(tmpl_run)
        for t in r.findall(qn('w:t')):
            r.remove(t)
        for tag in ('w:br','w:tab','w:cr'):
            for el in r.findall(qn(tag)):
                r.remove(el)
        t = r.makeelement(qn('w:t'), {qn('xml:space'):'preserve'})
        t.text = text
        r.append(t)
        new_p.append(r)
    else:
        np.add_run(text)
    return np

# ---- 1. COVER: name + student id (para 4) ----
set_paragraph_text(paras[4], "學生姓名:PoKai Huang      學號:26254793")

# ---- B1. Q3 confidence interval (append to para 25) ----
append_text(paras[25],
    "進一步以區間估計檢視,平均評分的 95% 信賴區間約為 [0.626, 0.660];此區間涵蓋檢定值 0.65,"
    "與上述 p = 0.431、無法拒絕原假設的結論一致。")

# ---- B2. Q2 variable-type classification (append to para 22) ----
append_text(paras[22],
    "就變數型態而言,level 屬定性(類別)變數;year 本身為數值,但本報告將其視為類別並以虛擬變數處理,"
    "而 price、num_subscribers、num_reviews、num_lectures、content_duration 與 rating 則均為定量變數。")

# ---- C2. Q10 clarify zero-price log handling (replace para 56) ----
set_paragraph_text(paras[56],
    "依公司經濟學家的建議,將問題五(含控制變數)的模型改以對數形式估計:因變數改為 log(訂閱人數)、"
    "價格自變數改為 log(價格)(資料集已提供對應的對數欄位;其中價格為零的免費課程,於取對數前先加上一微小常數,"
    "以避免 log(0) 無定義)。估計結果如表 5。")

# ---- B3. Fix Kaggle reference (replace para 71) ----
set_paragraph_text(paras[71],
    "andrewmvd. (2020). Udemy courses [Data set]. Kaggle. "
    "https://www.kaggle.com/datasets/andrewmvd/udemy-courses")

# ---- C1. Limitations paragraph: insert before 參考文獻 heading (para 70) ----
# template = a normal body paragraph (para 69 = conclusion body)
ref_para = paras[70]      # 參考文獻 heading
template  = paras[69]     # body style
insert_para_before(ref_para, "（補充)本分析仍有若干限制宜一併說明。資料屬單一時點的觀察性橫斷面,"
    "只能描述變數間的關聯、無法確立因果;免費課程對價量關係造成明顯干擾,使全樣本與付費樣本的結論在方向上出現反轉;"
    "加上模型僅能解釋約一到兩成的訂閱變異,課程品質、講師聲譽與行銷曝光等關鍵因素並未納入。因此上述結論較適合作為"
    "公司進入市場前的初步參考,而非精確的需求預測。", template)

doc.save(SRC)

# ---- verification ----
chk = Document(SRC)
out = []
for i, p in enumerate(chk.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    out.append(f"{i}\t{t}")
open("_verify.txt", "w", encoding="utf-8").write("\n".join(out))
print("SAVED_OK paras=%d" % len(chk.paragraphs))
